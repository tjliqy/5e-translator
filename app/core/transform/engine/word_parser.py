from app.core.transform.bean.bestiary import Bestiary, Attr, is_camp, is_group, BasicBean
import re
import os
import sys
import fitz  # PyMuPDF库，用于更精确地提取文本和字体信息
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

if sys.version_info < (3, 9):
    from typing import List
else:
    List = list

import os
import re
from typing import List, Dict, Tuple
import os
import re
import sys
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

if sys.version_info < (3, 9):
    from typing import List, Dict, Tuple
else:
    List = list
    Dict = dict
    Tuple = tuple

# 导入所需的库
try:
    import docx2txt
    import olefile
    import subprocess
    import zipfile
    import shutil
    import tempfile
    # 添加python-docx库用于提取字体信息
    try:
        import docx
    except ImportError:
        print("请安装python-docx库以支持精确的字体信息提取: pip install python-docx")
        docx = None
except ImportError:
    print("请安装必要的库: pip install docx2txt olefile")
    sys.exit(1)

class WordParser:
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.chinese_punctuation = r"。！？；…"
        self.chinese_sentence_pattern = re.compile(f'([^{self.chinese_punctuation}]*[{self.chinese_punctuation}])')
        self.paragraphs = []  # 存储解析后的段落
        self.paragraphs_with_format = []  # 存储带格式信息的段落

    def parse(self):
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"Word文件不存在: {self.file_path}")
        
        # 支持.doc和.docx格式
        file_ext = os.path.splitext(self.file_path)[1].lower()
        if file_ext not in ['.doc', '.docx']:
            raise ValueError(f"文件不是Word格式: {self.file_path}")
        
        try:
            self._parse_word_file()
            return self.paragraphs
        except Exception as e:
            raise RuntimeError(f"处理Word文件时出错: {str(e)}")
    
    def _try_parse_with_docx2txt(self):
        """尝试使用docx2txt解析文件"""
        try:
            print(f"尝试使用docx2txt解析文件: {self.file_path}")
            return docx2txt.process(self.file_path)
        except zipfile.BadZipFile:
            # 对于.doc文件，docx2txt可能会误判为zip文件
            raise RuntimeError("docx2txt不支持此格式的文件")
    
    def _try_parse_with_olefile_improved(self):
        """使用改进的olefile方法解析.doc文件"""
        print(f"尝试使用改进的olefile解析文件: {self.file_path}")
        
        if not olefile.isOleFile(self.file_path):
            raise ValueError(f"不是有效的OLE文件: {self.file_path}")
        
        with olefile.OleFileIO(self.file_path) as ole:
            # 打印所有可用的流，帮助调试
            print(f"OLE文件流列表: {ole.listdir()}")
            
            # 尝试寻找不同可能的文档流
            possible_streams = [
                # 标准Word文档流路径
                ['WordDocument'],
                # 可能的其他路径
                ['WordDocument', '1Table'],
                ['Macros'],
                ['SummaryInformation'],
                ['DocumentSummaryInformation']
            ]
            
            # 尝试所有可能的流
            for stream_path in possible_streams:
                try:
                    if ole.exists(stream_path):
                        print(f"找到流: {stream_path}")
                        doc_data = ole.openstream(stream_path).read()
                        text = self._extract_text_from_doc_data(doc_data)
                        if text and len(text.strip()) > 10:
                            return text
                except Exception as e:
                    print(f"尝试流 {stream_path} 失败: {str(e)}")
            
            # 尝试读取所有找到的流
            all_streams = ole.listdir()
            for stream in all_streams:
                try:
                    doc_data = ole.openstream(stream).read()
                    text = self._extract_text_from_doc_data(doc_data)
                    if text and len(text.strip()) > 10:
                        print(f"从流 {stream} 成功提取文本")
                        return text
                except Exception as e:
                    continue
            
            # 所有尝试都失败
            raise RuntimeError("无法从任何OLE流中提取有意义的文本")
    
    def _extract_text_from_doc_data(self, doc_data):
        """从doc二进制数据中提取文本"""
        try:
            # 优先尝试中文编码，添加更多编码选项
            encodings = ['gb2312', 'utf-8', 'latin-1', 'cp1252','gb18030', 'gbk', ]
            text = ""
            best_text = ""
            best_score = 0
            
            # 统计中文字符的正则表达式
            chinese_char_pattern = re.compile(r'[\u4e00-\u9fff]+')
            
            for encoding in encodings:
                try:
                    # 使用当前编码解码
                    text = doc_data.decode(encoding, errors='replace')
                    
                    # 过滤掉控制字符，保留可打印字符和换行符、制表符
                    filtered_text = ''.join([c if c.isprintable() or c in '\n\r\t' else ' ' for c in text])
                    
                    # 清理多余的空白字符，但保留段落结构（两个或多个换行符）
                    # 先处理多个空格
                    filtered_text = re.sub(r'[ \t]+', ' ', filtered_text)
                    # 保留段落结构（两个或多个换行符）
                    filtered_text = re.sub(r'(\n{2,})', '\n\n', filtered_text)
                    # 处理单个换行符（可能是行尾而非段落分隔）
                    filtered_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', filtered_text)
                    
                    # 去除首尾空白
                    filtered_text = filtered_text.strip()
                    
                    # 检查是否包含中文字符，作为评分标准
                    chinese_chars = chinese_char_pattern.findall(filtered_text)
                    chinese_text_length = sum(len(char) for char in chinese_chars)
                    
                    # 评分：文本长度 + 中文字符长度 * 2（优先选择包含更多中文的结果）
                    score = len(filtered_text) + chinese_text_length * 2
                    
                    # 更新最佳结果
                    if score > best_score:
                        best_score = score
                        best_text = filtered_text
                        print(f"尝试编码 {encoding} - 提取文本长度: {len(filtered_text)}, 中文字符数: {chinese_text_length}")
                        # 如果文本足够长且包含中文，直接返回
                        if len(filtered_text) > 100 and len(filtered_text) - chinese_text_length < 10:
                            print(f"使用编码 {encoding} 成功提取中文文本")
                            return filtered_text
                except Exception as e:
                    print(f"编码 {encoding} 解码失败: {str(e)}")
                    continue
            
            # 返回找到的最佳文本结果
            if best_text and len(best_text) > 10:
                print(f"最终选择的编码结果 - 文本长度: {len(best_text)}")
                return best_text
            
            # 如果所有编码都失败，返回空字符串
            return ""
        except Exception as e:
            print(f"文本提取过程出错: {str(e)}")
            import traceback
            traceback.print_exc()
            return ""  # 提取失败时返回空字符串
    
    def _try_parse_with_command_line_tool(self):
        """尝试使用系统命令行工具解析.doc文件"""
        print(f"尝试使用系统命令行工具解析文件: {self.file_path}")
        
        tools = [
            ['catdoc'],
            ['antiword'],
            ['libreoffice', '--headless', '--convert-to', 'txt', '--outdir', '{temp_dir}', '{file_path}'],
            ['soffice', '--headless', '--convert-to', 'txt', '--outdir', '{temp_dir}', '{file_path}']
        ]
        
        for tool_cmd in tools:
            try:
                if tool_cmd[0] in ['libreoffice', 'soffice']:
                    # 使用临时目录进行转换
                    with tempfile.TemporaryDirectory() as temp_dir:
                        cmd = [t.format(temp_dir=temp_dir, file_path=self.file_path) for t in tool_cmd]
                        subprocess.run(cmd, capture_output=True, check=True)
                        
                        # 查找生成的txt文件
                        base_name = os.path.splitext(os.path.basename(self.file_path))[0]
                        txt_file = os.path.join(temp_dir, base_name + '.txt')
                        if os.path.exists(txt_file):
                            with open(txt_file, 'r', encoding='utf-8', errors='replace') as f:
                                return f.read()
                else:
                    # 直接调用命令行工具
                    result = subprocess.run(tool_cmd + [self.file_path], capture_output=True, text=True, check=True)
                    return result.stdout
            except (subprocess.SubprocessError, FileNotFoundError):
                print(f"命令行工具 {tool_cmd[0]} 不可用或执行失败")
                continue
        
        # 所有命令行工具都失败
        return ""
    
    def _try_parse_as_text(self):
        """作为最后的尝试，直接将文件视为文本文件读取"""
        print(f"尝试直接以文本方式读取文件: {self.file_path}")
        try:
            # 尝试使用不同的编码读取
            encodings = ['utf-8', 'latin-1', 'cp1252', 'gbk', 'gb2312']
            
            for encoding in encodings:
                try:
                    with open(self.file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                        if text and len(text.strip()) > 10:
                            return text
                except Exception:
                    continue
        except Exception as e:
            print(f"直接读取文件失败: {str(e)}")
        
        return ""
    
    def _try_parse_with_docx(self):
        """使用python-docx库解析.docx文件，提取文本和字体信息"""
        if docx is None:
            raise ImportError("python-docx库未安装")
        
        try:
            doc = docx.Document(self.file_path)
            paragraphs_with_format = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # 获取段落的主要字体大小
                max_font_size = 0
                has_bold = False
                
                for run in para.runs:
                    if run.font.size:
                        font_size = run.font.size.pt
                        if font_size > max_font_size:
                            max_font_size = font_size
                    if run.bold:
                        has_bold = True
                
                # 如果没有明确的字体大小，使用默认值
                if max_font_size == 0:
                    max_font_size = 12  # 默认字体大小
                
                paragraphs_with_format.append({
                    'text': para.text,
                    'font_size': max_font_size,
                    'is_bold': has_bold
                })
            
            return paragraphs_with_format
        except Exception as e:
            print(f"使用python-docx解析失败: {str(e)}")
            raise RuntimeError("python-docx解析失败")

    def _parse_word_file(self):
        """解析Word文件，根据不同格式选择合适的解析方法"""
        file_ext = os.path.splitext(self.file_path)[1].lower()
        text = ""
        
        # 对于.docx文件，优先尝试使用python-docx获取字体信息
        if file_ext == '.docx' and docx is not None:
            try:
                self.paragraphs_with_format = self._try_parse_with_docx()
                # 构建纯文本内容
                text = '\n'.join([p['text'] for p in self.paragraphs_with_format])
            except Exception as e:
                print(f"使用python-docx获取字体信息失败: {str(e)}")
                # 失败后回退到原有的解析方法
                pass
        
        # 如果没有获取到带格式的信息，使用原有的解析方法
        if not self.paragraphs_with_format:
            # 尝试多种解析方法，按优先级尝试
            parsing_methods = [
                self._try_parse_with_docx2txt,
                self._try_parse_with_olefile_improved,
                self._try_parse_with_command_line_tool,
                self._try_parse_as_text
            ]
            
            for method in parsing_methods:
                try:
                    text = method()
                    if text and len(text.strip()) > 10:  # 确保提取到有意义的文本
                        break
                except Exception as e:
                    print(f"解析方法 {method.__name__} 失败: {str(e)}")
        
        if not text or len(text.strip()) <= 10:
            raise RuntimeError("所有解析方法均失败，无法提取有意义的文本内容")
        
        # 解析文本为段落
        self._parse_text_into_paragraphs(text)
    
    def _parse_text_into_paragraphs(self, text):
        """将文本解析为段落结构"""
        # 基于换行符和空行分段
        raw_paragraphs = text.split('\n\n')  # 连续两个换行符表示段落分隔
        
        # 如果分割结果不理想，尝试使用单个换行符分割
        if len(raw_paragraphs) < 3:
            raw_paragraphs = text.split('\n')
        
        # 过滤掉空段落并去除首尾空白
        raw_paragraphs = [p.strip() for p in raw_paragraphs if p.strip()]
        
        # 初始化当前段落
        current_paragraph = {
            'title': '',
            'content': []
        }
        
        # 标题识别逻辑
        if self.paragraphs_with_format:
            # 使用字体大小进行标题识别
            self._parse_with_font_size_info()
        else:
            # 使用传统的文本特征识别
            for i, para in enumerate(raw_paragraphs):
                # 按照文本特征识别标题：
                is_potential_title = len(para) < 100 and \
                                    (para.isupper() or para.istitle() or para[0].isupper())
                has_following_content = i < len(raw_paragraphs) - 1 and len(raw_paragraphs[i+1]) > 150
                
                if is_potential_title and has_following_content:
                    # 如果当前有正在处理的段落，先保存它
                    if current_paragraph['title'] or current_paragraph['content']:
                        self.paragraphs.append(current_paragraph)
                    
                    # 开始新的段落，设置标题
                    current_paragraph = {
                        'title': para,
                        'content': []
                    }
                else:
                    # 正文内容，添加到当前段落
                    current_paragraph['content'].append(para)
            
            # 添加最后一个段落（如果有）
            if current_paragraph['title'] or current_paragraph['content']:
                self.paragraphs.append(current_paragraph)

    def _parse_with_font_size_info(self):
        """使用字体大小信息识别标题并构建段落结构"""
        if not self.paragraphs_with_format:
            return
        
        # 计算文档中字体大小的统计信息，以确定标题阈值
        font_sizes = [p['font_size'] for p in self.paragraphs_with_format]
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
        # 标题通常比正文大1.2-1.5倍
        title_font_threshold = avg_font_size * 1.2
        
        current_paragraph = {
            'title': '',
            'content': []
        }
        
        # 查找最大的字体大小，通常是主标题
        max_font_size = max(font_sizes) if font_sizes else 12
        
        for i, para_info in enumerate(self.paragraphs_with_format):
            text = para_info['text'].strip()
            font_size = para_info['font_size']
            is_bold = para_info['is_bold']
            
            # 判断是否为标题的条件
            # 1. 字体明显大于平均大小或为文档中最大的字体
            # 2. 文本长度适中（通常标题不会太长）
            # 3. 可能是粗体
            is_title = (
                (font_size >= title_font_threshold or font_size == max_font_size) and 
                len(text) < 200 and 
                (is_bold or len(text) < 100)
            )
            
            # 对于中文文档的特殊处理
            if re.search(r'[\u4e00-\u9fff]', text):
                # 中文字符较多，标题通常更简洁
                is_title = is_title and len(text) < 150
            
            if is_title:
                # 如果当前有正在处理的段落，先保存它
                if current_paragraph['title'] or current_paragraph['content']:
                    self.paragraphs.append(current_paragraph)
                
                # 开始新的段落，设置标题
                current_paragraph = {
                    'title': text,
                    'content': []
                }
            else:
                # 正文内容，添加到当前段落
                current_paragraph['content'].append(text)
        
        # 添加最后一个段落（如果有）
        if current_paragraph['title'] or current_paragraph['content']:
            self.paragraphs.append(current_paragraph)

    def __split_sentences(self, text: str) -> List[str]:
        """
        将文本分割为适当大小的块
        """
        # 知识库中单段文本长度
        CHUNK_SIZE = 150
        
        # 知识库中相邻文本重合长度
        OVERLAP_SIZE = 30
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=OVERLAP_SIZE
        )
        
        return text_splitter.split_text(text)
    
    def get_documents(self) -> (List[Document]):
        """
        将解析后的段落转换为Document对象列表
        """
        documents = []        
        for p in self.paragraphs:
            # 处理标题
            if p['title']:
                documents.append(Document(
                    page_content=p['title'], 
                    metadata={'title': p['title'], 'source': self.file_path}
                ))
            
            # 处理内容
            for c in p['content']:
                # 分割长文本
                split_contents = self.__split_sentences(c)
                for content_chunk in split_contents:
                    documents.append(Document(
                        page_content=content_chunk, 
                        metadata={'title': p['title'] if p['title'] else '无标题段落', 'source': self.file_path}
                    ))
        
        return documents